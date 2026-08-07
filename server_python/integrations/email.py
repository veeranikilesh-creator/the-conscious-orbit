"""Email service — sends report DOCX to users after admin approval.
Degrades gracefully: without SMTP config, logs instead of sending.
"""
import io
import os
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication


def _slug(name: str) -> str:
    s = (name or "venture").lower()
    return "".join(c if c.isalnum() or c == "-" else "-" for c in s).strip("-") or "venture"


def build_report_docx(report_data: dict, client_data: dict | None = None) -> io.BytesIO:
    """Render the admin-approved report as a real Office Open XML .docx.

    Returns a BytesIO positioned at 0 so callers can .read() or attach it
    directly. Never mutates the inputs.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page margins — a bit tighter than default for a denser report look.
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    # Title
    title = doc.add_heading(report_data.get("name", "Strategy Report"), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        f'{(report_data.get("vertical") or "startups").upper()}  ·  PUBLISHED'
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x8C, 0x6D, 0x58)

    # Divider
    doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Score panel
    score = report_data.get("adminScore") if report_data.get("adminScore") is not None else report_data.get("score", 0)
    verdict = report_data.get("adminVerdict") or (
        "GO" if (report_data.get("decision") == 1) else "PIVOT" if report_data.get("decision") == 0 else "PENDING"
    )

    score_p = doc.add_paragraph()
    score_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = score_p.add_run(f"Conscious Orbital Score: {score}/100")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0xD4, 0xAF, 0x37)

    verdict_p = doc.add_paragraph()
    verdict_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vr = verdict_p.add_run(f"Verdict: {verdict}")
    vr.bold = True
    vr.font.size = Pt(12)

    # Company & contact table
    doc.add_heading("Company", level=1)
    client = client_data or {}
    if isinstance(client_data, dict):
        table = doc.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"
        for label, value in [
            ("Company", client.get("company") or report_data.get("name", "—")),
            ("Industry", client.get("industry") or report_data.get("vertical") or "—"),
            ("Stage", client.get("stage") or "—"),
            ("Geography", client.get("geography") or "—"),
            ("Business Model", client.get("businessModel") or "—"),
            ("Contact", client.get("contact") or "—"),
            ("Email", client.get("email") or "—"),
        ]:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(value)
            for para in row[0].paragraphs:
                for rr in para.runs:
                    rr.bold = True

    # Admin analysis
    analysis = (report_data.get("adminAnalysis") or "").strip()
    if analysis:
        doc.add_heading("Our Analysis", level=1)
        doc.add_paragraph(analysis)

    # Strengths
    strengths = (report_data.get("adminStrengths") or "").strip()
    if strengths:
        doc.add_heading("Strengths", level=1)
        for line in [ln.strip() for ln in strengths.splitlines() if ln.strip()]:
            doc.add_paragraph(line, style="List Bullet")

    # Risks
    risks = (report_data.get("adminRisks") or "").strip()
    if risks:
        doc.add_heading("Risks & Concerns", level=1)
        for line in [ln.strip() for ln in risks.splitlines() if ln.strip()]:
            doc.add_paragraph(line, style="List Bullet")

    # Admin note
    note = (report_data.get("approvalNote") or "").strip()
    if note:
        doc.add_heading("Note from Our Team", level=1)
        doc.add_paragraph(note)

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("The Conscious Orbit")
    fr.italic = True
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x9A, 0x9A, 0x9A)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _build_report_html(report_data: dict) -> str:
    """Generate the report HTML with admin's manual analysis."""
    score = report_data.get("score") or report_data.get("adminScore") or 0
    verdict = report_data.get("adminVerdict") or report_data.get("decision") or "PENDING"
    verdict_map = {"GO": "GO", "CONDITIONAL": "CONDITIONAL", "PIVOT": "PIVOT", "REJECT": "REJECT"}
    verdict_label = verdict_map.get(verdict, verdict)
    verdict_color = (
        "#16a34a" if verdict == "GO"
        else "#dc2626" if verdict == "REJECT"
        else "#ea580c" if verdict == "PIVOT"
        else "#ca8a04"
    )

    analysis = report_data.get("adminAnalysis", "")
    analysis_section = ""
    if analysis:
        analysis_section = f'''
        <div style="background: #f0fdf4; border: 2px solid #16a34a; border-radius: 12px; padding: 20px; margin: 20px 0;">
            <h3 style="margin: 0 0 12px 0; color: #16a34a;">Our Analysis</h3>
            <p style="color: #166534; line-height: 1.6; white-space: pre-wrap;">{analysis}</p>
        </div>'''

    strengths = report_data.get("adminStrengths", "")
    strengths_section = ""
    if strengths:
        items = "".join(f"<li style='margin-bottom:4px;'>{s.strip()}</li>" for s in strengths.split("\n") if s.strip())
        strengths_section = f'''
        <div style="background: #ecfdf5; border-radius: 12px; padding: 20px; margin: 20px 0;">
            <h3 style="margin: 0 0 8px 0; color: #059669;">Strengths</h3>
            <ul style="color: #065f46; margin: 0; padding-left: 20px;">{items}</ul>
        </div>'''

    risks = report_data.get("adminRisks", "")
    risks_section = ""
    if risks:
        items = "".join(f"<li style='margin-bottom:4px;'>{r.strip()}</li>" for r in risks.split("\n") if r.strip())
        risks_section = f'''
        <div style="background: #fef2f2; border-radius: 12px; padding: 20px; margin: 20px 0;">
            <h3 style="margin: 0 0 8px 0; color: #dc2626;">Risks & Concerns</h3>
            <ul style="color: #991b1b; margin: 0; padding-left: 20px;">{items}</ul>
        </div>'''

    note = report_data.get("approvalNote", "")
    note_section = ""
    if note:
        note_section = f'''
        <div style="background: #eff6ff; border-radius: 12px; padding: 20px; margin: 20px 0;">
            <h3 style="margin: 0 0 8px 0; color: #2563eb;">Note from Our Team</h3>
            <p style="color: #1e40af; line-height: 1.6;">{note}</p>
        </div>'''

    return f"""<!DOCTYPE html>
<html>
<head><meta charset="utf-8"><title>{report_data.get("name", "Report")} — Strategy Report</title></head>
<body style="font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px;">
  <h1 style="color: #1a1a1a;">{report_data.get("name", "Report")}</h1>
  <p style="color: #666;">{(report_data.get("vertical", "startups")).upper()} · PUBLISHED</p>
  <div style="background: #f8f9fa; border: 2px solid #D4AF37; border-radius: 12px; padding: 20px; margin: 20px 0;">
    <h2 style="margin: 0; color: #D4AF37;">Conscious Orbital Score: {score}/100</h2>
    <span style="display: inline-block; padding: 4px 12px; border-radius: 20px; background: {verdict_color}; color: white; font-weight: bold; margin-top: 8px;">{verdict_label}</span>
  </div>
  {analysis_section}
  {strengths_section}
  {risks_section}
  {note_section}
  <p style="color: #999; font-size: 12px; margin-top: 40px;">The Conscious Orbit</p>
</body>
</html>"""


def send_report_email(report_data: dict, client_data: dict) -> dict:
    """Send the report as a .doc email attachment to the user.

    Returns {"sent": bool, "message": str}.
    """
    if isinstance(client_data, str):
        recipient_email = None
    else:
        recipient_email = (client_data or {}).get("email") if isinstance(client_data, dict) else getattr(client_data, "email", None)

    if not recipient_email:
        return {"sent": False, "message": "No email address on file for this client."}

    smtp_host = os.getenv("SMTP_HOST", "")
    if not smtp_host:
        print(f"[EMAIL] SMTP not configured. Would send report \"{report_data.get('name')}\" to {recipient_email}")
        return {"sent": False, "message": "SMTP not configured — email logged but not sent."}

    html = _build_report_html(report_data)
    slug = _slug(report_data.get("name"))

    msg = MIMEMultipart()
    msg["From"] = os.getenv("SMTP_FROM", '"The Conscious Orbit" <reports@consciousorbit.com>')
    msg["To"] = recipient_email
    msg["Subject"] = f"Your Venture Strategy Report: {report_data.get('name', 'Report')}"

    msg.attach(MIMEText(html, "html"))

    # Real Office Open XML .docx \u2014 python-docx renders a proper Word document
    # that any Office/Google Docs/LibreOffice client can open cleanly.
    try:
        docx_buf = build_report_docx(report_data, client_data if isinstance(client_data, dict) else None)
        docx_attachment = MIMEApplication(
            docx_buf.read(),
            _subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        docx_attachment.add_header(
            "Content-Disposition", "attachment", filename=f"{slug}-strategy-report.docx"
        )
        msg.attach(docx_attachment)
    except Exception as e:
        # If python-docx isn't installed or rendering fails, fall back to the
        # HTML-as-.doc attachment so the email still ships with something.
        print(f"[EMAIL] .docx render failed ({e}); attaching HTML .doc fallback.")
        doc_attachment = MIMEApplication(("\ufeff" + html).encode("utf-8"), _subtype="msword")
        doc_attachment.add_header("Content-Disposition", "attachment", filename=f"{slug}-strategy-report.doc")
        msg.attach(doc_attachment)

    try:
        port = int(os.getenv("SMTP_PORT", "587"))
        secure = os.getenv("SMTP_SECURE", "false").lower() == "true"
        with smtplib.SMTP(smtp_host, port) as server:
            if secure:
                server.starttls()
            smtp_user = os.getenv("SMTP_USER", "")
            smtp_pass = os.getenv("SMTP_PASS", "")
            if smtp_user:
                server.login(smtp_user, smtp_pass)
            server.send_message(msg)
        return {"sent": True, "message": f"Report emailed to {recipient_email}"}
    except Exception as e:
        print(f"[EMAIL] Failed to send to {recipient_email}: {e}")
        return {"sent": False, "message": f"Email failed: {e}"}
